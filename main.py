from dotenv import load_dotenv
import os
import requests
from langchain_openai import ChatOpenAI
from loadCSV import load_tickets_from_csv
from langchain.tools import tool
from langchain.agents import initialize_agent, AgentType
from loadJira import load_tickets_by_project_and_fix_version as load_tickets_from_jira
import json
from docx import Document

# Load environment variables
load_dotenv()

# Initialize LLM
llm = ChatOpenAI(
    api_key=os.getenv("OPENAI_API_KEY"),
    model="gpt-4o-mini"
)

# === Core logic functions (not tools) ===

def summarize_ticket(title, description, key=None):
    prompt = f"""
            You are writing internal release notes.

            Return the response in JSON with the following fields:
            - key: the Jira ticket ID (use '{key}' if available)
            - title: the ticket title (use exact title provided)
            - release_note: a clear, professional sentence describing what the ticket is about, client's facing text
            - key_features: optional, a list of bullet points of key features or highlights

            Title: {title}
            Description: {description}
            Ticket Key: {key or 'UNKNOWN'}
            """
    response = llm.invoke(prompt)
    import json
    try:
        return json.loads(response.content.strip())
    except:
        return {
            "key": key or "UNKNOWN",
            "title": title,
            "summary": response.content.strip(),
            "key_features": []
        }


def categorize_ticket(issue_type):
    if issue_type.lower() == "story":
        return "Features"
    elif issue_type.lower() == "bug":
        return "Bug Fixes"
    elif issue_type.lower() == "task":
        return "Improvements"
    else:
        return "Other"


def generate_release_notes(csv_path, version):
    tickets = load_tickets_from_csv(csv_path, version)

    categorized_notes = {
        "Features": [],
        "Bug Fixes": [],
        "Improvements": [],
        "Other": []
    }

    for t in tickets:
        summary = t.get("Summary", "")
        description = t.get("Description", "")
        issue_type = t.get("Issue Type", "")
        key = t.get("Key", "UNKNOWN")

        note_dict = summarize_ticket(summary, description, key)
        category = categorize_ticket(issue_type)
        categorized_notes[category].append(note_dict)

    return categorized_notes


def format_release_notes_md(notes_dict, version):
    lines = [f"# Release Notes - Version {version}", ""]

    for category, items in notes_dict.items():
        if not items:
            continue
        lines.append(f"## {category}")
        for item in items:
                lines.append(f"- **{item.get('key', '')}**: {item.get('title', '')}")
                lines.append(f"  - Summary: {item.get('summary', '')}")
                if item.get("key_features"):
                    for feat in item["key_features"]:
                        lines.append(f"    - {feat}")

        lines.append("")

    return "\n".join(lines)

def format_release_notes_docx(notes_dict, version, output_file="release_notes.docx"):
    doc = Document()
    doc.add_heading(f'Release Notes - Version {version}', 0)

    for category, entries in notes_dict.items():
        if not entries:
            continuepytho
        doc.add_heading(category, level=1)

        for entry in entries:
            doc.add_heading(entry["key"], level=2)
            doc.add_paragraph(f"Title: {entry.get('title', '')}")
            doc.add_paragraph(f"Summary: {entry.get('summary', '')}")
            if entry.get("key_features"):
                doc.add_paragraph("Key Features:")
                for feat in entry["key_features"]:
                    doc.add_paragraph(feat, style='List Bullet')


    doc.save(output_file)
    print(f"✅ DOCX saved to {output_file}")

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])


# === Tool-wrapped functions for LangChain ===

@tool
def summarize_ticket_tool(prompt: str) -> str:
    """
    Summarize a Jira ticket.
    Input format: 'summary=...; description=...'
    """
    try:
        parts = dict(item.strip().split("=", 1) for item in prompt.split(";"))
        summary = parts.get("summary", "")
        description = parts.get("description", "")
    except Exception as e:
        return f"Invalid input format. Use: summary=...; description=...\nError: {str(e)}"

    return json.dumps(summarize_ticket(summary, description))


@tool
def categorize_ticket_tool(issue_type_str: str) -> str:
    """
    Categorize Jira ticket by type. Input: 'story', 'bug', etc.
    """
    return categorize_ticket(issue_type_str.strip())


@tool
def generate_release_notes_wrapper(query: str) -> str:
    """
    Generate formatted release notes.
    Input format: 'csv=...; version=...'
    """
    try:
        parts = dict(item.strip().split("=", 1) for item in query.split(";"))
        csv_path = parts.get("csv", "tickets.csv")
        version = parts.get("version", "v1.0.0")
    except Exception as e:
        return f"Failed to parse input. Use: 'csv=...; version=...'\nError: {str(e)}"

    notes_dict = generate_release_notes(csv_path, version)
    return format_release_notes_md(notes_dict, version)

@tool
def generate_docx_release_notes(query: str) -> str:
    """
    Generate a .docx file with formatted release notes.
    Input: 'csv=...; version=...'
    """
    try:
        parts = dict(item.strip().split("=", 1) for item in query.split(";"))
        csv_path = parts.get("csv", "tickets.csv")
        version = parts.get("version", "v1.0.0")
    except Exception as e:
        return f"Invalid input: {str(e)}"

    notes = generate_release_notes(csv_path, version)
    filename = f"release_notes_{version}.docx"
    format_release_notes_docx(notes, version, output_file=filename)
    return f"✅ DOCX file created: {filename}"


# === Register tools and initialize agent ===

tools = [
    summarize_ticket_tool,
    categorize_ticket_tool,
    generate_release_notes_wrapper,
    generate_docx_release_notes
]

agent = initialize_agent(
    tools=tools,
    llm=llm,
    agent=AgentType.ZERO_SHOT_REACT_DESCRIPTION,
    verbose=True
)

# === Manual execution for testing ===

def get_tickets_and_version():
    source = input("Load tickets from CSV or Jira? (csv/jira): ").strip().lower()

    if source == "csv":
        csv_path = input("Enter CSV path (default: tickets.csv): ").strip() or "tickets.csv"
        version = input("Enter release version (e.g. v1.0.0): ").strip()
        tickets = load_tickets_from_csv(csv_path, version)
    elif source == "jira":
        project_key = input("Enter project key (e.g. TNG4): ").strip()
        fix_versions_input = input("Enter Fix Versions (comma-separated): ").strip()
        fix_versions = [v.strip() for v in fix_versions_input.split(",") if v.strip()]
        version = ", ".join(fix_versions)

        # Call your already formatted Jira loader
        from loadJira import load_tickets_by_project_and_fix_version
        email = os.getenv("JIRA_EMAIL")
        api_token = os.getenv("JIRA_API_TOKEN")
        base_url = os.getenv("JIRA_BASE_URL")

        tickets = load_tickets_by_project_and_fix_version(
                                                                project_key,
                                                                fix_versions,
                                                                email,
                                                                api_token,
                                                                base_url
                                                            )

    else:
        raise ValueError("Unknown source. Use 'csv' or 'jira'.")

    return tickets, version


if __name__ == "__main__":
    print("📝 Manual test — generate release notes\n")

    tickets, version = get_tickets_and_version()

    categorized_notes = {
        "Features": [],
        "Bug Fixes": [],
        "Improvements": [],
        "Other": []
    }

    for t in tickets:
        summary = t.get("Summary", "") or t.get("summary", "")
        description = t.get("Description", "") or t.get("description", "")
        issue_type = t.get("Issue Type", "") or t.get("issue_type", "")
        key = t.get("Key", "") or t.get("key", "UNKNOWN")

        note_dict = summarize_ticket(summary, description, key)
        category = categorize_ticket(issue_type)
        categorized_notes[category].append(note_dict)

    markdown = format_release_notes_md(categorized_notes, version)
    print(markdown)

    format_release_notes_docx(categorized_notes, version, f"release_notes_{version}.docx")
    # 🔽 Extract readable text
    docx_path = f"release_notes_{version}.docx"
    readable_text = extract_text_from_docx(docx_path)
    print("\n📄 Extracted DOCX text:\n")
    print(readable_text)

    # Optionally save to a .txt file
    with open(f"release_notes_{version}.txt", "w", encoding="utf-8") as f:
        f.write(readable_text)
        print(f"✅ Plain text saved to release_notes_{version}.txt")



